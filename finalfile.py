import requests
import json
import os
from docx import Document
from datetime import datetime
import shutil
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors


#            *** To fetch the data form cyclops ***
def get_json_data():
    cyclops_template = os.environ.get('TEMPLATE')
    CTMS_URL = os.environ.get('CTMS_URL')
    S3 =os.environ.get('S3')
    if S3 == "red":
     cyclops_url = "https://cyclopsui.imedidata.com/"
    else:
      cyclops_url = "https://cyclopsui-sandbox.imedidata.net"
    url = f"{cyclops_url}/api/v0/url-configurations/CTMS/{cyclops_template}/{CTMS_URL}"
    payload = {}
    headers = {
        'Access-Token': os.environ.get('cyclops'), #stored in secure variables, will be adding in secrete manager after testing
        'Content-Type': 'application/json'
    }
    response = requests.get(url, headers=headers, data=payload)
    if response.status_code == 200:
        json_dict = response.json()
        return json_dict
    else:
        print(f"Error: {response.status_code} - {response.text}")
        return None


#             *** To update the PIR file ***
def replace_word_file_content():
    json_data = get_json_data()
    current_directory = os.getcwd()
    now = datetime.now()
    deploy_date= now.strftime("%d-%m-%Y")
    print("Current Working Directory:", current_directory)
    file_name="CTMS_PIR.docx"
    word_file_path = os.path.join(current_directory, file_name)
    print("word_file_path:" , word_file_path)
    replacements = {
        "Product_Version": json_data["global"]["CTMS_VERSION"],
        "Environment_URL": os.environ.get('CTMS_URL') ,
        "Git_Branch": json_data["global"]["GIT_BRANCH"],
        "Deploy_By": json_data["global"]["deploy_by"],
        "Deploy_Date":deploy_date
    }
    document = Document(file_name)
    print("Before Replacement:")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)
                
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for target_word, replacement_word in replacements.items():
                    cell.text = cell.text.replace(target_word, replacement_word)

    print("\nAfter Replacement:")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)

    # Save the updated Word document
    output_path = "PIR.docx"
    document.save(output_path)
    print(f"\nUpdated Word file saved to: {output_path}")


#             *** To convert the PIR word into pdf format ***
def apply_styles(pdf_table, table_index):
    if table_index == 0:
        style = TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('RIGHTPADDING', (0, 0), (-1, -1), -200)

        ])
    elif table_index == 1:
        style = TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ])
    else:
        # Add more conditions for additional tables if needed
        style = TableStyle([])

    pdf_table.setStyle(style)

def count_tables(docx_path):
    doc = Document(docx_path)
    table_count = 0
    # To create PDF file
    pdfname=os.environ.get('CTMS_URL')
    pdffilename = f"{pdfname}_PIR.pdf"
    pdf = SimpleDocTemplate(pdffilename, pagesize=letter)
    pdf_tables = []
    for table_index, table in enumerate(doc.tables):
        table_count += 1   
        # Extracting table data
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        # To Create PDF table
        pdf_table = Table(table_data)
        # Applying style based on the table index
        apply_styles(pdf_table, table_index)
        # Add the styled table to the list
        pdf_tables.append(pdf_table)
    # Build the PDF file with all tables
    pdf.build(pdf_tables)


#    *** To copy PIR.pdf to the reports directory ***
def copyfile():
    url=os.environ.get('CTMS_URL')
    pirfilename = f"{url}_PIR.pdf" 
    destination_folder = "./reports/"
    destination_path = os.path.join(destination_folder, pirfilename)
    shutil.copy(pirfilename, destination_path)


#   *** To upload IR and PIR to the jira ticket
def upload_attachment():
    access_token = os.environ.get('Access_Token')
    ticket_id = os.environ.get('Jira_Ticket_Id')
    url=os.environ.get('CTMS_URL')
    json_data = get_json_data()
    Git_Branch = json_data["global"]["GIT_BRANCH"]
    pirfilename = f"{url}_PIR.pdf"
    # irfilename = f"{url}_{Git_Branch}.pdf" 
    url = f"https://jira.mdsol.com/rest/api/2/issue/{ticket_id}/attachments"
    script_directory = os.path.dirname(os.path.abspath(__file__))
    # file_path = os.path.join(script_directory, f"./reports/{irfilename}")
    file_path = os.path.join(script_directory, f"./reports/{pirfilename}")

    headers = {
    'X-Atlassian-Token': 'nocheck',
    'Authorization': f'Bearer {access_token}'
    }
    files = {"file": open(file_path, "rb")}
    response = requests.post(url, headers=headers, files=files)
    print(response.status_code)
    print(response.text)


#execution starts from here
replace_word_file_content()                     # 1. To call replacement function
wordfile_to_pdffile = count_tables('PIR.docx')  # 2. To convert wordfile to pdf file
copyfile()                                      # 3. To copy PIR.pdf to the reports directory
upload_attachment()                             # 4. To upload IR and PIR to the jira ticket

