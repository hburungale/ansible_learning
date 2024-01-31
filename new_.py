from docx import Document
import os


def replace_and_print_word_file_content(file_path, replacements):

    document = Document(file_path)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for target_word, replacement_word in replacements.items():
                    cell.text = cell.text.replace(target_word, replacement_word)

    print("\nAfter Replacement:")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)

    # Save the updated Word document
    output_path = "updated.docx"
    document.save(output_path)
    print(f"\nUpdated Word file saved to: {output_path}")

if __name__ == "__main__":
   
    current_directory = os.getcwd()
    print("Current Working Directory:", current_directory)
    file_name="app_codes\CTMS_Val3_HDC PIR.docx"
    word_file_path = os.path.join(current_directory, file_name)
    print("word_file_path:" , word_file_path)
    replacements = {
        "Product_Name": os.environ.get('Product_Name'),
        "Product_Version": os.environ.get('Product_Version'),
        "Environment_URL": os.environ.get('Environment_URL'),
        "Git_Branch": os.environ.get('Git_Branch'),
        "Deploy_By": os.environ.get('Deploy_By'),
        "Deploy_Date":os.environ.get('Deploy_Date')
    }

    # Call the function to replace and print the content
    replace_and_print_word_file_content(word_file_path, replacements)
