from docx import Document
import os
# import cyclops


def replace_and_print_word_file_content(file_path, replacements):

    document = Document(file_path)
    # for paragraph in document.paragraphs:
    #     print(paragraph.text)

    # print("Before Replacement:")
    # for table in document.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             print(cell.text)

    # Replace target words with replacement words
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for target_word, replacement_word in replacements.items():
                    cell.text = cell.text.replace(target_word, replacement_word)

    print("\nAfter Replacement:")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)

    # Save the updated Word document
    output_path = "Final_PIR.docx"
    document.save(output_path)
    print(f"\nUpdated Word file saved to: {output_path}")

if __name__ == "__main__":
   
    current_directory = os.getcwd()
    print("Current Working Directory:", current_directory)
    file_name="app_codes\CTMS_Val1_PIR.docx"
    word_file_path = os.path.join(current_directory, file_name)
    print("word_file_path:" , word_file_path)
    word_file_path = "/Users/hburungale/Desktop/Hrithik/IR_automation/ansible_learning/CTMS_PIR_copy.docx"

    replacements = {
        "Product_Name": "CTMS",
        "Product_Version": "2202",
        "Environment_URL": "ctms.net",
        "Git_Branch": "123456",
        "Deploy_By": "hburungale",
        "Deploy_Date":"04-01-2024"
    }

    # Call the function to replace and print the content
    replace_and_print_word_file_content(word_file_path, replacements)
