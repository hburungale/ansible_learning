from docx import Document
import os
from datetime import datetime
from FetchCyclopsData import get_json_data

def replace_and_print_word_file_content(file_path, replacements):
    document = Document(file_path)
    print("Before Replacement:")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)
                
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for target_word, replacement_word in replacements.items():
                    cell.text = cell.text.replace(target_word, replacement_word)

    print("\nAfter Replacement:")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                print(cell.text)

    # Save the updated Word document
    output_path = "Distro_PIR.docx"
    document.save(output_path)
    print(f"\nUpdated Word file saved to: {output_path}")

if __name__ == "__main__":
    json_data = get_json_data()
    current_directory = os.getcwd()
    now = datetime.now()
    deploy_date= now.strftime("%d-%m-%Y")
    print("Current Working Directory:", current_directory)
    file_name="CTMS_PIR.docx"
    word_file_path = os.path.join(current_directory, file_name)
    print("word_file_path:" , word_file_path)
 

    replacements = {
        "Product_Version": json_data["global"]["CTMS_VERSION"],
        "Environment_URL": os.environ.get('Environment_URL') ,
        "Git_Branch": json_data["global"]["GIT_BRANCH"],
        "Deploy_By": json_data["global"]["deploy_by"],
        "Deploy_Date":deploy_date
    }
    

    # Call the function to replace and print the content
    replace_and_print_word_file_content(word_file_path, replacements)
