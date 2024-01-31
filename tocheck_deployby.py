from docx import Document
import os

# def replace_table_cell_value(table, target, replacement):
#     for row in table.rows:
#         for cell in row.cells:
#             if target in cell.text:
#                 cell.text = cell.text.replace(target, replacement)

def replace_and_print_word_file_content(file_path, replacements):
    document = Document(file_path)

    for paragraph in document.paragraphs:
        print(paragraph.text)

    # print("\nBefore Replacement:")
    # for table in document.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             print(cell.text)

    # Replace target words in the table cells
    # for table in document.tables:
    #     replace_table_cell_value(table, 'Deploy_By', replacements.get('Deploy_By', ''))
    #     replace_table_cell_value(table, 'Deploy_Date', replacements.get('Deploy_Date', ''))

    # print("\nAfter Replacement:")
    # for table in document.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             print(cell.text)

    # Save the updated Word document
    output_path = "newfile.docx"
    document.save(output_path)
    print(f"\nUpdated Word file saved to: {output_path}")

if __name__ == "__main__":
    # current_directory = os.getcwd()
    file_name = "/Users/hburungale/Desktop/Hrithik/IR_automation/ansible_learning/CTMS_Val3_HDC PIR.docx"
    # word_file_path = os.path.join(current_directory, file_name)

    # Specify the replacements as a dictionary
    replacements = {
        "Deploy_By": "HB",
        "Deploy_Date": "jan3",
    }

    # Call the function to replace and print the content
    replace_and_print_word_file_content(file_name, replacements)
